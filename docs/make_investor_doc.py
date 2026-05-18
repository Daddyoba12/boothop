from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
MID_NAVY  = RGBColor(0x16, 0x21, 0x3E)
BLUE      = RGBColor(0x0F, 0x3D, 0x66)
GREY      = RGBColor(0x55, 0x55, 0x55)
RED       = RGBColor(0x99, 0x00, 0x00)

def set_heading(para, level, text):
    sizes = {1: 20, 2: 15, 3: 12}
    colors = {1: DARK_NAVY, 2: MID_NAVY, 3: BLUE}
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, DARK_NAVY)
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10 if level == 2 else 8)
    para.paragraph_format.space_after = Pt(4)

def add_rule(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_inline(para, text):
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            para.add_run(part)

def add_body_para(doc, text, space_before=2, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    parse_inline(para, text)
    return para

def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    parse_inline(para, text)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table_from_lines(doc, lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\|[-| :]+\|$', stripped):
            continue
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if any(cells):
            rows.append(cells)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j in range(num_cols):
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = row.cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)
            if i == 0:
                shade_cell(cell, 'DCE8F8')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

# ── Cover page ──────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(80)
r = cover.add_run('BootHop')
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = DARK_NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(6)
r2 = sub.add_run('Investor Pack  ·  May 2026')
r2.font.size = Pt(14)
r2.font.color.rgb = GREY

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf.paragraph_format.space_before = Pt(10)
r3 = conf.add_run('Confidential — Not for Distribution')
r3.italic = True
r3.font.size = Pt(11)
r3.font.color.rgb = RED

site = doc.add_paragraph()
site.alignment = WD_ALIGN_PARAGRAPH.CENTER
site.paragraph_format.space_before = Pt(60)
r4 = site.add_run('boothop.co.uk')
r4.font.size = Pt(12)
r4.font.color.rgb = BLUE

doc.add_page_break()

# ── Parse markdown ───────────────────────────────────────────────────────────
with open(r'C:\Users\babso\Desktop\boothop\boothop\docs\INVESTOR_PACK.md', 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

table_buf = []
i = 0

while i < len(lines):
    line = lines[i]

    # Collect table rows
    if line.strip().startswith('|'):
        table_buf.append(line)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_lines(doc, table_buf)
            table_buf = []

    # H1
    if re.match(r'^# [^#]', line):
        text = line[2:].strip()
        if 'Investor Pack' not in text:
            para = doc.add_paragraph()
            set_heading(para, 1, text)
        i += 1; continue

    # H2
    if re.match(r'^## ', line):
        text = line[3:].strip()
        if not text.startswith('Confidential'):
            para = doc.add_paragraph()
            set_heading(para, 2, text)
        i += 1; continue

    # H3
    if re.match(r'^### ', line):
        para = doc.add_paragraph()
        set_heading(para, 3, line[4:].strip())
        i += 1; continue

    # HR
    if line.strip() == '---':
        add_rule(doc)
        i += 1; continue

    # Bullet
    if re.match(r'^- ', line):
        add_bullet(doc, line[2:].strip())
        i += 1; continue

    # Code block
    if line.strip().startswith('```'):
        i += 1
        code = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code.append(lines[i])
            i += 1
        i += 1
        para = doc.add_paragraph()
        run = para.add_run('\n'.join(code))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        continue

    # Empty line
    if line.strip() == '':
        i += 1; continue

    # Body paragraph
    add_body_para(doc, line.strip())
    i += 1

if table_buf:
    add_table_from_lines(doc, table_buf)

out = r'C:\Users\babso\Desktop\boothop\boothop\docs\BootHop_Investor_Pack.docx'
doc.save(out)
print('Saved:', out)
